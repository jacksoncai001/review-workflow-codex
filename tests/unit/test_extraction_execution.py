from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from review_workflow.adapters.corpus import SourceRecord, compute_source_id
from review_workflow.adapters.extraction import (
    DoclingExtractor,
    MarkItDownExtractor,
    PypdfExtractor,
    ToolAvailability,
)
from review_workflow.adapters.filesystem import sha256_file
from review_workflow.application.corpus_service import CorpusService
from review_workflow.domain.models import ExecutionProfile


def test_pypdf_extractor_keeps_page_locator_even_for_a_blank_page(tmp_path: Path) -> None:
    source = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with source.open("wb") as stream:
        writer.write(stream)
    digest = sha256_file(source)

    result = PypdfExtractor().extract(
        source,
        source_id=compute_source_id(digest),
        source_hash=digest,
    )

    assert result.parser == "pypdf"
    assert result.page_count == 1
    assert result.units[0].page == 1
    assert "page_has_no_extractable_text" in result.warnings


def test_full_profile_persists_grobid_tei_beside_primary_extraction(tmp_path: Path) -> None:
    source = tmp_path / "article.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with source.open("wb") as stream:
        writer.write(stream)
    digest = sha256_file(source)
    record = SourceRecord(
        source_id=compute_source_id(digest),
        source_hash=digest,
        original_path=source,
        size_bytes=source.stat().st_size,
        extension=".pdf",
    )

    class FakeGrobid:
        def process_fulltext(self, source_path):
            assert source_path == source
            return type("Document", (), {"tei_xml": "<TEI><text/></TEI>"})()

    report = CorpusService(
        availability=ToolAvailability(pypdf=True, grobid=True),
        grobid_adapter=FakeGrobid(),
    ).extract_and_index(
        workspace=tmp_path / "workspace",
        sources=[record],
        profile=ExecutionProfile.FULL,
    )

    tei_path = Path(report["extractions"][0]["tei_path"])
    assert tei_path.read_text(encoding="utf-8") == "<TEI><text/></TEI>\n"


def test_markitdown_adapter_extracts_docx_without_writing_beside_source(tmp_path: Path) -> None:
    source = tmp_path / "draft.docx"
    document = Document()
    document.add_heading("Review boundary", level=1)
    document.add_paragraph("A reusable local extraction.")
    document.save(source)
    before = set(tmp_path.iterdir())
    digest = sha256_file(source)

    result = MarkItDownExtractor().extract(
        source,
        source_id=compute_source_id(digest),
        source_hash=digest,
    )

    assert "Review boundary" in result.markdown
    assert "reusable local extraction" in result.markdown
    assert set(tmp_path.iterdir()) == before


def test_docling_adapter_can_process_local_pdf_when_layout_extra_is_installed(
    tmp_path: Path,
) -> None:
    pytest.importorskip("docling")
    source = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with source.open("wb") as stream:
        writer.write(stream)
    before = set(tmp_path.iterdir())
    digest = sha256_file(source)

    result = DoclingExtractor().extract(
        source,
        source_id=compute_source_id(digest),
        source_hash=digest,
    )

    assert result.parser == "docling"
    assert result.extraction_id.startswith("ext-")
    assert set(tmp_path.iterdir()) == before
