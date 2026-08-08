from __future__ import annotations

import json
from pathlib import Path

import httpx
from docx import Document

from review_workflow.adapters.discovery import OaLocation, WorkRecord
from review_workflow.adapters.extraction import ToolAvailability
from review_workflow.application.corpus_service import CorpusService
from review_workflow.application.service import WorkflowService

CODEX_QUESTIONS = [
    "Who is the primary reader?",
    "What decision should the review support?",
    "Which nearest review must this paper differ from?",
]


def write_docx(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Synthetic review input", level=1)
    document.add_paragraph(text)
    document.save(path)


def write_valid_pdf(path: Path, title: str) -> bytes:
    from pypdf import PdfWriter

    path.parent.mkdir(parents=True, exist_ok=True)
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_metadata({"/Title": title})
    with path.open("wb") as stream:
        writer.write(stream)
    with path.open("ab") as stream:
        stream.write(b"\n% synthetic-padding\n" + b"x" * 2048)
    return path.read_bytes()


def write_and_register(
    service: WorkflowService,
    workspace: Path,
    *,
    artifact_id: str,
    kind: str,
    relative_path: str,
    phase: str,
    dependencies: list[str],
) -> None:
    path = workspace / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps({"artifact": artifact_id}) + "\n", encoding="utf-8")
    service.artifact_register(
        workspace=workspace,
        artifact_id=artifact_id,
        kind=kind,
        relative_path=relative_path,
        producer="goal-path-test",
        phase=phase,
        dependencies=dependencies,
    )


class DiscoveryClient:
    def search(self, query: str) -> list[WorkRecord]:
        return [
            WorkRecord(
                work_id="anchor",
                doi="10.1000/anchor",
                title=f"Nearest technical review for {query}",
                abstract="Review boundary and diagnostic methods",
                is_review=True,
            ),
            WorkRecord(
                work_id="contrast",
                doi="10.1000/contrast",
                title="Limitations and conflicting diagnostic evidence",
                abstract="A comparison that challenges proxy equivalence",
            ),
            WorkRecord(
                work_id="bridge",
                doi="10.1000/bridge",
                title="Validation methods for external measurements",
                abstract="Experimental methods connect measurements to diagnosis",
            ),
        ]


class OpenAccessClient:
    def lookup(self, doi: str) -> OaLocation:
        if doi == "10.1000/contrast":
            return OaLocation(doi=doi, is_oa=False, oa_status="closed")
        return OaLocation(
            doi=doi,
            is_oa=True,
            oa_status="gold",
            landing_url=f"https://example.org/{doi.rsplit('/', 1)[-1]}",
            pdf_url=f"https://example.org/{doi.rsplit('/', 1)[-1]}.pdf",
            version="publishedVersion",
            license="cc-by",
        )


class AllOpenAccessClient:
    def lookup(self, doi: str) -> OaLocation:
        slug = doi.rsplit("/", 1)[-1]
        return OaLocation(
            doi=doi,
            is_oa=True,
            oa_status="gold",
            landing_url=f"https://example.org/{slug}",
            pdf_url=f"https://example.org/{slug}.pdf",
            version="publishedVersion",
            license="cc-by",
        )


def test_initial_recommendations_wait_for_reading_before_phase_1(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        CorpusService,
        "detect_tools",
        staticmethod(
            lambda: ToolAvailability(
                markitdown=False,
                docling=False,
                grobid=False,
                pypdf=True,
                python_docx=True,
            )
        ),
    )
    initial = (tmp_path / "operator-input" / "draft.docx").resolve()
    write_docx(initial, "Initial evidence supplied by the operator.")
    workspace = (tmp_path / "review-workspace").resolve()
    service = WorkflowService()
    service.project_init(
        workspace=workspace,
        project_id="initial-recommendations",
        review_type="technical",
        execution_profile="windows-lite",
        publication_mode="single",
    )
    service.phase_next(workspace=workspace, target="0")
    service.source_inventory(
        workspace=workspace,
        input_paths=[str(initial)],
        run_extraction=False,
    )
    pdf_by_url = {
        "https://example.org/anchor.pdf": write_valid_pdf(
            tmp_path / "initial-anchor.pdf", "Nearest technical review"
        ),
        "https://example.org/contrast.pdf": write_valid_pdf(
            tmp_path / "initial-contrast.pdf", "Limitations and conflicting evidence"
        ),
    }

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            headers={"Content-Type": "application/pdf"},
            content=pdf_by_url[str(request.url)],
            request=request,
        )

    discovery = service.literature_discover(
        workspace=workspace,
        search_lanes=["external measurement validation limits"],
        count=2,
        clients=[DiscoveryClient()],
        unpaywall_client=AllOpenAccessClient(),
        download_http_client=httpx.Client(transport=httpx.MockTransport(handler)),
    )

    assert len(discovery["automatic_downloads"]) == 2
    waiting = service.project_status(workspace)
    assert waiting["phase"] == "0"
    assert waiting["status"] == "WAITING_USER"
    reading_packet = service.question_packet_get(workspace)
    assert reading_packet["packet_type"] == "recommended_reading"
    acknowledged = service.recommended_reading_acknowledge(
        workspace=workspace,
        packet_id=reading_packet["packet_id"],
        reading_notes=[
            "The anchor defines the nearest scope.",
            "The contrast limits proxy equivalence.",
        ],
    )

    assert acknowledged["phase"] == "0"
    assert acknowledged["status"] == "RUNNING"
    assert acknowledged["step"] == "initial_recommendations_read"
    service.phase_next(workspace=workspace, target="1")
    inventory = service.source_inventory(
        workspace=workspace,
        input_paths=[str(initial), str(workspace / "acquisitions")],
        run_extraction=True,
    )
    assert inventory["corpus_counts"]["sources"] == 3


def test_goal_path_expands_literature_resumes_scoping_and_completes_phase_7(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        CorpusService,
        "detect_tools",
        staticmethod(
            lambda: ToolAvailability(
                markitdown=False,
                docling=False,
                grobid=False,
                pypdf=True,
                python_docx=True,
            )
        ),
    )
    initial = (tmp_path / "operator-input" / "draft.docx").resolve()
    write_docx(initial, "Initial evidence about an external measurement chain.")
    workspace = (tmp_path / "review-workspace").resolve()
    service = WorkflowService()
    service.project_init(
        workspace=workspace,
        project_id="literature-refresh",
        review_type="technical",
        execution_profile="windows-lite",
        publication_mode="single",
    )
    service.phase_next(workspace=workspace, target="0")
    service.phase_next(workspace=workspace, target="1")
    initial_inventory = service.source_inventory(
        workspace=workspace,
        input_paths=[str(initial)],
        run_extraction=True,
    )
    assert initial_inventory["corpus_counts"]["sources"] == 1
    service.phase_next(workspace=workspace, target="2A")

    packet = service.question_packet_open(workspace, CODEX_QUESTIONS)
    answer = service.answer_packet_record(
        workspace=workspace,
        packet_id=packet["packet_id"],
        answer_payload={
            "operator_answers": ["Engineers", "Choose a method", "Review X"],
            "operator_questions": [
                "What validation evidence is missing?",
                "Which boundary is still uncertain?",
                "What literature would change the outline?",
            ],
            "codex_answers": [
                "Direct comparison",
                "Proxy equivalence",
                "A validation bridge",
            ],
            "changed_assumptions": ["Validation must be explicit"],
            "unresolved_tensions": ["Proxy equivalence"],
            "new_search_lanes": ["external measurement validation limits"],
        },
    )

    returned = service.project_status(workspace)
    assert answer["literature_refresh"]["status"] == "search_required"
    assert returned["phase"] == "0"
    assert returned["step"] == "literature_refresh:scope-round-1"
    assert returned["resume_action"]["step"] == "round_1_after_literature_refresh"

    pdf_by_url = {
        "https://example.org/anchor.pdf": write_valid_pdf(
            tmp_path / "anchor-source.pdf", "Nearest technical review"
        ),
        "https://example.org/bridge.pdf": write_valid_pdf(
            tmp_path / "bridge-source.pdf", "Validation methods"
        ),
    }

    def download_handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(
            200,
            headers={"Content-Type": "application/pdf"},
            content=pdf_by_url[str(request.url)],
            request=request,
        )

    discovery = service.literature_discover(
        workspace=workspace,
        search_lanes=["external measurement validation limits"],
        count=3,
        clients=[DiscoveryClient()],
        unpaywall_client=OpenAccessClient(),
        auto_download_open=True,
        download_http_client=httpx.Client(transport=httpx.MockTransport(download_handler)),
    )

    assert len(discovery["recommendations"]) == 3
    assert not discovery["automatic_download_errors"], discovery
    assert len(discovery["automatic_downloads"]) == 2
    requests = {item["doi"]: item for item in service.acquisition_request_list(workspace)}
    assert requests["10.1000/anchor"]["status"] == "fulfilled"
    assert requests["10.1000/bridge"]["status"] == "fulfilled"
    assert requests["10.1000/contrast"]["status"] == "open"
    waiting = service.project_status(workspace)
    assert waiting["status"] == "WAITING_ACQUISITION"
    assert waiting["phase"] == "0"

    supplied = workspace / "acquisitions" / "user-supplied" / "contrast.pdf"
    write_valid_pdf(supplied, "Limitations and conflicting diagnostic evidence")
    imported = service.acquisition_import(
        workspace=workspace,
        request_id=requests["10.1000/contrast"]["request_id"],
        relative_path=str(supplied.relative_to(workspace)),
        expected_identity={
            "doi": "10.1000/contrast",
            "title": "Limitations and conflicting diagnostic evidence",
        },
        download_metadata={
            "content_type": "application/pdf",
            "source_url": "https://doi.org/10.1000/contrast",
            "access_basis": "publisher_permitted",
            "license_or_terms": "operator attestation",
            "document_role": "full_text",
            "observed_doi": "10.1000/contrast",
            "observed_title": "Limitations and conflicting diagnostic evidence",
        },
    )

    reading_wait = imported["workflow_resume"]
    assert reading_wait["phase"] == "0"
    assert reading_wait["status"] == "WAITING_USER"
    assert imported["workflow_resume"]["corpus_inventory"]["corpus_counts"]["sources"] == 4
    manifest_lines = (
        (workspace / "inputs" / "manifest.jsonl").read_text(encoding="utf-8").splitlines()
    )
    assert len(manifest_lines) == 4
    assert initial.is_file()

    reading_packet = service.question_packet_get(workspace)
    assert reading_packet["packet_type"] == "recommended_reading"
    resumed = service.recommended_reading_acknowledge(
        workspace=workspace,
        packet_id=reading_packet["packet_id"],
        reading_notes=[
            "The anchor narrows the nearest-review boundary.",
            "The contrast challenges proxy equivalence.",
            "The bridge clarifies validation requirements.",
        ],
    )
    assert resumed["status"] == "RUNNING"
    assert resumed["phase"] == "2A"
    assert resumed["step"] == "round_1_after_literature_refresh"
    assert resumed["resume_action"] is None

    for round_number in (2, 3):
        packet = service.question_packet_open(workspace, CODEX_QUESTIONS)
        service.answer_packet_record(
            workspace=workspace,
            packet_id=packet["packet_id"],
            answer_payload={
                "operator_answers": [
                    f"reader answer {round_number}",
                    f"decision answer {round_number}",
                    f"nearest review answer {round_number}",
                ],
                "operator_questions": [
                    "Which contradiction now matters most?",
                    "What claim must remain qualified?",
                    "Can the evidence support this positioning?",
                ],
                "codex_answers": [
                    "The validation boundary",
                    "The proxy-equivalence claim",
                    "Yes, with explicit limitations",
                ],
                "changed_assumptions": [],
                "unresolved_tensions": [] if round_number == 3 else ["scope still open"],
                "new_search_lanes": [],
            },
        )

    write_and_register(
        service,
        workspace,
        artifact_id="scope-brief",
        kind="scope_brief",
        relative_path="phases/phase-2/scope-brief.json",
        phase="2A",
        dependencies=["source-manifest"],
    )
    write_and_register(
        service,
        workspace,
        artifact_id="architecture-scorecard",
        kind="architecture_scorecard",
        relative_path="phases/phase-2/architecture-scorecard.json",
        phase="2E",
        dependencies=["scope-brief"],
    )
    for phase in ("2B", "2C", "2D", "2E"):
        service.phase_next(workspace=workspace, target=phase)
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_2e_outline",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="3")
    write_and_register(
        service,
        workspace,
        artifact_id="claim-matrix",
        kind="claim_matrix",
        relative_path="phases/phase-3/claim-matrix.json",
        phase="3",
        dependencies=["architecture-scorecard"],
    )
    service.phase_next(workspace=workspace, target="4")
    write_and_register(
        service,
        workspace,
        artifact_id="manuscript",
        kind="manuscript",
        relative_path="manuscripts/review.md",
        phase="4",
        dependencies=["claim-matrix"],
    )
    service.phase_next(workspace=workspace, target="5")
    write_and_register(
        service,
        workspace,
        artifact_id="citation-audit",
        kind="citation_audit",
        relative_path="phases/phase-5/citation-audit.json",
        phase="5",
        dependencies=["manuscript", "claim-matrix"],
    )
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_5_citation",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="6")
    write_and_register(
        service,
        workspace,
        artifact_id="review-report",
        kind="review_report",
        relative_path="phases/phase-6/review-report.json",
        phase="6",
        dependencies=["citation-audit"],
    )
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_6_review",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="7")
    bundle = service.reproducibility_bundle_create(workspace)
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_7_final",
        approved_by="operator",
    )
    completed = service.workflow_complete(workspace)

    assert completed["status"] == "COMPLETED"
    assert completed["loop_counters"]["mutual_scoping_rounds"] == 3
    assert Path(bundle["bundle_path"]).is_file()
