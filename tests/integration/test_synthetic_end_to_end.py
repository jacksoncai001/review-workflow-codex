from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from review_workflow.application.service import WorkflowService

QUESTIONS = [
    ["Who is the primary reader?", "What decision should they make?", "What is in scope?"],
    ["Which review is nearest?", "What is the sharpest difference?", "What is excluded?"],
    ["What remains unresolved?", "Is the evidence feasible?", "Single or companion paper?"],
]


def write_and_register(
    service: WorkflowService,
    workspace: Path,
    *,
    artifact_id: str,
    kind: str,
    relative_path: str,
    phase: str,
    dependencies: list[str],
    version: int = 1,
    replace_existing: bool = False,
) -> None:
    path = workspace / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps({"artifact": artifact_id, "version": version}) + "\n", encoding="utf-8"
    )
    service.artifact_register(
        workspace=workspace,
        artifact_id=artifact_id,
        kind=kind,
        relative_path=relative_path,
        producer="synthetic-test",
        phase=phase,
        dependencies=dependencies,
        replace_existing=replace_existing,
    )


def test_synthetic_review_runs_three_rounds_repairs_citation_and_packages_bundle(
    tmp_path: Path,
) -> None:
    source = (tmp_path / "operator-input" / "draft.docx").resolve()
    source.parent.mkdir()
    document = Document()
    document.add_heading("External measurement", level=1)
    document.add_paragraph("A measurement chain connects a signal to a bounded decision.")
    document.save(source)
    source_hash_before = source.read_bytes()
    workspace = (tmp_path / "review-workspace").resolve()
    service = WorkflowService()
    service.project_init(
        workspace=workspace,
        project_id="synthetic-review",
        review_type="technical",
        execution_profile="windows-lite",
        publication_mode="single",
    )
    service.phase_next(workspace=workspace, target="0")
    service.phase_next(workspace=workspace, target="1")
    inventory = service.source_inventory(
        workspace=workspace,
        input_paths=[str(source)],
        run_extraction=True,
    )
    assert inventory["corpus_counts"]["chunks"] > 0
    service.phase_next(workspace=workspace, target="2A")

    for round_number, questions in enumerate(QUESTIONS, start=1):
        packet = service.question_packet_open(workspace, questions)
        service.answer_packet_record(
            workspace=workspace,
            packet_id=packet["packet_id"],
            answer_payload={
                "operator_answers": [f"operator answer {round_number}-{i}" for i in range(3)],
                "operator_questions": [f"Operator question {round_number}-{i}?" for i in range(3)],
                "codex_answers": [f"codex answer {round_number}-{i}" for i in range(3)],
                "changed_assumptions": [],
                "unresolved_tensions": [] if round_number == 3 else ["scope still open"],
                "new_search_lanes": [],
            },
        )

    write_and_register(
        service,
        workspace,
        artifact_id="scope-brief",
        kind="scope_brief",
        relative_path="phases/phase-2/scope-brief.json",
        phase="2A",
        dependencies=[],
    )
    write_and_register(
        service,
        workspace,
        artifact_id="architecture-scorecard",
        kind="architecture_scorecard",
        relative_path="phases/phase-2/architecture-scorecard.json",
        phase="2E",
        dependencies=["scope-brief"],
    )
    for phase in ("2B", "2C", "2D", "2E"):
        service.phase_next(workspace=workspace, target=phase)
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_2e_outline",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="3")
    write_and_register(
        service,
        workspace,
        artifact_id="claim-matrix",
        kind="claim_matrix",
        relative_path="phases/phase-3/claim-matrix.json",
        phase="3",
        dependencies=["architecture-scorecard"],
    )
    service.phase_next(workspace=workspace, target="4")
    write_and_register(
        service,
        workspace,
        artifact_id="manuscript",
        kind="manuscript",
        relative_path="manuscripts/review.md",
        phase="4",
        dependencies=["claim-matrix"],
    )
    service.phase_next(workspace=workspace, target="5")
    write_and_register(
        service,
        workspace,
        artifact_id="citation-audit",
        kind="citation_audit",
        relative_path="phases/phase-5/citation-audit.json",
        phase="5",
        dependencies=["manuscript", "claim-matrix"],
    )

    service.return_route(
        workspace=workspace,
        failure_payload={
            "failure_id": "wrong-direction-claim",
            "kind": "citation_semantic_mismatch",
            "reason": "One source contradicts rather than supports the sentence",
            "repair_phase": "3",
            "origin_phase": "5",
            "origin_step": "citation_gate",
            "invalidated_artifacts": ["claim-matrix"],
            "preserved_artifacts": [],
        },
    )
    for artifact_id, kind, path, phase, dependencies in (
        (
            "claim-matrix",
            "claim_matrix",
            "phases/phase-3/claim-matrix.json",
            "3",
            ["architecture-scorecard"],
        ),
        ("manuscript", "manuscript", "manuscripts/review.md", "4", ["claim-matrix"]),
        (
            "citation-audit",
            "citation_audit",
            "phases/phase-5/citation-audit.json",
            "5",
            ["manuscript", "claim-matrix"],
        ),
    ):
        write_and_register(
            service,
            workspace,
            artifact_id=artifact_id,
            kind=kind,
            relative_path=path,
            phase=phase,
            dependencies=dependencies,
            version=2,
            replace_existing=True,
        )
    service.return_resume(
        workspace=workspace,
        failure_id="wrong-direction-claim",
        resolution_note=(
            "The claim, manuscript sentence, and citation audit now preserve the source's "
            "contradicting direction."
        ),
        evidence_artifact_ids=["claim-matrix", "manuscript", "citation-audit"],
    )
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_5_citation",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="6")
    write_and_register(
        service,
        workspace,
        artifact_id="review-report",
        kind="review_report",
        relative_path="phases/phase-6/review-report.json",
        phase="6",
        dependencies=["citation-audit"],
    )
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_6_review",
        approved_by="operator",
    )
    service.phase_next(workspace=workspace, target="7")
    bundle = service.reproducibility_bundle_create(workspace)
    service.gate_approve(
        workspace=workspace,
        gate_id="phase_7_final",
        approved_by="operator",
    )
    completed = service.workflow_complete(workspace)

    assert completed["status"] == "COMPLETED"
    assert completed["loop_counters"]["mutual_scoping_rounds"] == 3
    assert Path(bundle["bundle_path"]).is_file()
    assert set(bundle["required_kinds"]) <= set(bundle["present_kinds"])
    assert source.read_bytes() == source_hash_before
