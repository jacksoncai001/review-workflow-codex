from __future__ import annotations

from pathlib import Path

from docx import Document

from review_workflow.application.service import WorkflowService
from review_workflow.domain.phases import PhaseId


def test_inventory_can_extract_docx_into_reusable_local_corpus_without_touching_source(
    tmp_path: Path,
) -> None:
    source = (tmp_path / "operator-input" / "draft.docx").resolve()
    source.parent.mkdir()
    document = Document()
    document.add_heading("Measurement chain", level=1)
    document.add_paragraph("A sensor measures an external signal without opening the device.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Method"
    table.cell(0, 1).text = "Evidence"
    table.cell(1, 0).text = "Imaging"
    table.cell(1, 1).text = "Direct observation"
    document.save(source)
    source_before = source.read_bytes()

    workspace = (tmp_path / "review-workspace").resolve()
    service = WorkflowService()
    service.project_init(
        workspace=workspace,
        project_id="synthetic-review",
        review_type="technical",
        execution_profile="windows-lite",
    )
    service.phase_next(workspace=workspace, target=PhaseId.PHASE_0.value)
    service.phase_next(workspace=workspace, target=PhaseId.PHASE_1.value)

    report = service.source_inventory(
        workspace=workspace,
        input_paths=[str(source.parent)],
        run_extraction=True,
    )
    hits = service.evidence_search(
        workspace=workspace,
        query="external signal",
        top_k=5,
    )

    assert len(report["sources"]) == 1
    assert report["extractions"][0]["parser"] == "python-docx"
    assert report["extractions"][0]["table_count"] == 1
    assert report["corpus_counts"]["chunks"] >= 2
    assert hits[0]["source_id"] == report["sources"][0]["source_id"]
    assert source.read_bytes() == source_before
    assert Path(report["extractions"][0]["markdown_path"]).is_relative_to(workspace)


def test_reextracting_same_source_and_configuration_reuses_immutable_extraction(
    tmp_path: Path,
) -> None:
    source = (tmp_path / "operator-input" / "draft.docx").resolve()
    source.parent.mkdir()
    document = Document()
    document.add_paragraph("Stable extracted material for reuse on another computer.")
    document.save(source)
    workspace = (tmp_path / "review-workspace").resolve()
    service = WorkflowService()
    service.project_init(
        workspace=workspace,
        project_id="synthetic-review",
        review_type="narrative",
        execution_profile="windows-lite",
    )

    first = service.source_inventory(
        workspace=workspace,
        input_paths=[str(source)],
        run_extraction=True,
    )
    second = service.source_inventory(
        workspace=workspace,
        input_paths=[str(source)],
        run_extraction=True,
    )

    assert first["extractions"][0]["extraction_id"] == second["extractions"][0]["extraction_id"]
    assert second["extractions"][0]["reused"] is True
    assert second["corpus_counts"] == first["corpus_counts"]
