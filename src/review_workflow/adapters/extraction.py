"""Execution-profile routing for versioned document extraction."""

from __future__ import annotations

import hashlib
import importlib.metadata
import json
from enum import StrEnum
from pathlib import Path
from typing import Any, Protocol

import yaml
from pydantic import Field

from review_workflow.domain.models import ExecutionProfile, StrictModel


class ExtractionTool(StrEnum):
    MARKITDOWN = "markitdown"
    DOCLING = "docling"
    GROBID = "grobid"
    PYPDF = "pypdf"
    PYTHON_DOCX = "python-docx"


class ToolAvailability(StrictModel):
    markitdown: bool = False
    docling: bool = False
    grobid: bool = False
    pypdf: bool = False
    python_docx: bool = False


class ExtractionUnit(StrictModel):
    """One locator-bearing unit of extracted text."""

    unit_id: str
    text: str
    page: int | None = Field(default=None, ge=1)
    section: str | None = None
    kind: str = "paragraph"


class ExtractionResult(StrictModel):
    """Parser-neutral, serializable extraction result."""

    schema_version: int = 1
    source_id: str
    source_hash: str = Field(pattern=r"^[0-9a-f]{64}$")
    extraction_id: str
    parser: str
    parser_version: str
    markdown: str
    units: list[ExtractionUnit] = Field(default_factory=list)
    page_count: int | None = Field(default=None, ge=0)
    figure_count: int = Field(default=0, ge=0)
    tables: list[list[list[str]]] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class Extractor(Protocol):
    """Interface implemented by optional local extraction adapters."""

    name: str

    def extract(
        self,
        source_path: Path,
        *,
        source_id: str,
        source_hash: str,
    ) -> ExtractionResult: ...


class ExtractionPlan(StrictModel):
    source_type: str
    profile: ExecutionProfile
    parsers: list[ExtractionTool] = Field(default_factory=list)
    degraded_capabilities: list[str] = Field(default_factory=list)
    mandatory_manual_checks: list[str] = Field(default_factory=list)


class DomainProfile(StrictModel):
    schema_version: int = 1
    profile_id: str = Field(pattern=r"^[a-z0-9][a-z0-9-]{1,62}[a-z0-9]$")
    display_name: str
    domains: list[str] = Field(min_length=1)
    synonyms: dict[str, list[str]] = Field(default_factory=dict)
    inclusion_clues: list[str] = Field(default_factory=list)
    exclusion_clues: list[str] = Field(default_factory=list)
    evidence_priorities: list[str] = Field(default_factory=list)
    measurement_chains: list[dict[str, Any]] = Field(default_factory=list)
    newcomer_concepts: list[str] = Field(min_length=1)
    prohibited_equivalences: list[str] = Field(default_factory=list)


class ExtractionRouter:
    """Choose available parsers without coupling core imports to parser packages."""

    def __init__(self, profile: ExecutionProfile, available_tools: ToolAvailability) -> None:
        self.profile = profile
        self.available = available_tools

    def route(self, source_path: Path) -> ExtractionPlan:
        extension = source_path.suffix.lower()
        if extension in {".docx", ".doc", ".rtf", ".odt"}:
            parsers: list[ExtractionTool] = []
            degraded: list[str] = []
            if extension == ".docx" and self.available.python_docx:
                parsers.append(ExtractionTool.PYTHON_DOCX)
            elif self.available.markitdown:
                parsers.append(ExtractionTool.MARKITDOWN)
            elif self.available.docling:
                parsers.append(ExtractionTool.DOCLING)
            else:
                degraded.append("office_document_extractor_unavailable")
            return ExtractionPlan(
                source_type="office_document",
                profile=self.profile,
                parsers=parsers,
                degraded_capabilities=degraded,
            )
        if extension != ".pdf":
            return ExtractionPlan(
                source_type=extension.lstrip(".") or "unknown",
                profile=self.profile,
                degraded_capabilities=["unsupported_source_type"],
            )

        parsers = []
        degraded = []
        manual = []
        if self.available.docling:
            parsers.append(ExtractionTool.DOCLING)
        elif self.available.pypdf:
            parsers.append(ExtractionTool.PYPDF)
            degraded.append("layout_figures_tables_may_be_incomplete")
        elif self.available.markitdown:
            parsers.append(ExtractionTool.MARKITDOWN)
            degraded.append("layout_figures_tables_may_be_incomplete")
        else:
            degraded.append("pdf_extractor_unavailable")

        if self.profile is ExecutionProfile.FULL and self.available.grobid:
            parsers.append(ExtractionTool.GROBID)
        else:
            degraded.append("grobid_unavailable")
            manual.append("bibliography_and_intext_citation_structure")

        if self.available.markitdown and ExtractionTool.MARKITDOWN not in parsers:
            parsers.append(ExtractionTool.MARKITDOWN)

        return ExtractionPlan(
            source_type="pdf",
            profile=self.profile,
            parsers=parsers,
            degraded_capabilities=degraded,
            mandatory_manual_checks=manual,
        )


def compute_extraction_id(
    *,
    source_hash: str,
    parser: str,
    parser_version: str,
    parser_config: dict[str, Any],
    schema_version: int,
) -> str:
    """Create a stable extraction ID from every output-affecting input."""
    payload = json.dumps(
        {
            "source_hash": source_hash,
            "parser": parser,
            "parser_version": parser_version,
            "parser_config": parser_config,
            "schema_version": schema_version,
        },
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )
    return f"ext-{hashlib.sha256(payload.encode('utf-8')).hexdigest()}"


def load_domain_profile(path: Path) -> DomainProfile:
    """Load a domain extension that cannot override engine or privacy policy."""
    payload = yaml.safe_load(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("Domain profile must contain a YAML object")
    return DomainProfile.model_validate(payload)


class PypdfExtractor:
    """Page-preserving PDF text extractor with a lazy optional import."""

    name = "pypdf"

    def extract(
        self,
        source_path: Path,
        *,
        source_id: str,
        source_hash: str,
    ) -> ExtractionResult:
        from pypdf import PdfReader

        version = _package_version("pypdf")
        extraction_id = compute_extraction_id(
            source_hash=source_hash,
            parser=self.name,
            parser_version=version,
            parser_config={"page_separator": "markdown_heading"},
            schema_version=1,
        )
        reader = PdfReader(source_path)
        units: list[ExtractionUnit] = []
        warnings: list[str] = []
        markdown_pages: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            units.append(
                ExtractionUnit(
                    unit_id=f"page-{page_number}",
                    text=text,
                    page=page_number,
                    kind="page",
                )
            )
            markdown_pages.append(f"## Page {page_number}\n\n{text}".rstrip())
            if not text:
                warnings.append("page_has_no_extractable_text")
        return ExtractionResult(
            source_id=source_id,
            source_hash=source_hash,
            extraction_id=extraction_id,
            parser=self.name,
            parser_version=version,
            markdown="\n\n".join(markdown_pages) + "\n",
            units=units,
            page_count=len(reader.pages),
            warnings=sorted(set(warnings)),
        )


class PythonDocxExtractor:
    """Structure-preserving DOCX baseline with headings and tables."""

    name = "python-docx"

    def extract(
        self,
        source_path: Path,
        *,
        source_id: str,
        source_hash: str,
    ) -> ExtractionResult:
        from docx import Document

        version = _package_version("python-docx")
        extraction_id = compute_extraction_id(
            source_hash=source_hash,
            parser=self.name,
            parser_version=version,
            parser_config={"headings": True, "tables": True, "inline_shapes": "count"},
            schema_version=1,
        )
        document = Document(source_path)
        units: list[ExtractionUnit] = []
        markdown_parts: list[str] = []
        current_section: str | None = None
        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading"):
                digits = "".join(character for character in style_name if character.isdigit())
                level = min(max(int(digits or "1"), 1), 6)
                current_section = text
                markdown_parts.append(f"{'#' * level} {text}")
                kind = "heading"
            else:
                markdown_parts.append(text)
                kind = "paragraph"
            units.append(
                ExtractionUnit(
                    unit_id=f"paragraph-{index}",
                    text=text,
                    section=current_section,
                    kind=kind,
                )
            )

        tables: list[list[list[str]]] = []
        for table_number, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            padded = [row + [""] * (width - len(row)) for row in rows]
            markdown_parts.append("| " + " | ".join(padded[0]) + " |")
            markdown_parts.append("| " + " | ".join(["---"] * width) + " |")
            for row in padded[1:]:
                markdown_parts.append("| " + " | ".join(row) + " |")
            table_text = "\n".join(" | ".join(row) for row in padded)
            units.append(
                ExtractionUnit(
                    unit_id=f"table-{table_number}",
                    text=table_text,
                    section=current_section,
                    kind="table",
                )
            )
        return ExtractionResult(
            source_id=source_id,
            source_hash=source_hash,
            extraction_id=extraction_id,
            parser=self.name,
            parser_version=version,
            markdown="\n\n".join(markdown_parts) + "\n",
            units=units,
            figure_count=len(document.inline_shapes),
            tables=tables,
        )


class MarkItDownExtractor:
    """Microsoft MarkItDown adapter loaded only when selected."""

    name = "markitdown"

    def extract(
        self,
        source_path: Path,
        *,
        source_id: str,
        source_hash: str,
    ) -> ExtractionResult:
        from markitdown import MarkItDown

        version = _package_version("markitdown")
        extraction_id = compute_extraction_id(
            source_hash=source_hash,
            parser=self.name,
            parser_version=version,
            parser_config={"plugins": False},
            schema_version=1,
        )
        converted = MarkItDown(enable_plugins=False).convert(str(source_path))
        text = str(converted.text_content or "").strip()
        warnings = [] if text else ["document_has_no_extractable_text"]
        units = [ExtractionUnit(unit_id="document-1", text=text, kind="document")]
        return ExtractionResult(
            source_id=source_id,
            source_hash=source_hash,
            extraction_id=extraction_id,
            parser=self.name,
            parser_version=version,
            markdown=text + ("\n" if text else ""),
            units=units,
            warnings=warnings,
        )


class DoclingExtractor:
    """Layout-aware Docling adapter loaded only when selected."""

    name = "docling"

    def extract(
        self,
        source_path: Path,
        *,
        source_id: str,
        source_hash: str,
    ) -> ExtractionResult:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import DocumentConverter, PdfFormatOption

        version = _package_version("docling")
        extraction_id = compute_extraction_id(
            source_hash=source_hash,
            parser=self.name,
            parser_version=version,
            parser_config={"export": "markdown"},
            schema_version=1,
        )
        if source_path.suffix.lower() == ".pdf":
            pipeline_options = PdfPipelineOptions()
            # torch.compile is optional acceleration. Disabling it avoids locale-dependent
            # template decoding failures on Chinese Windows while preserving inference.
            pipeline_options.layout_options.engine_options.compile_model = False
            converter = DocumentConverter(
                format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)}
            )
        else:
            converter = DocumentConverter()
        document = converter.convert(str(source_path)).document
        markdown = document.export_to_markdown().strip()
        units = [ExtractionUnit(unit_id="document-1", text=markdown, kind="document")]
        tables = [[] for _ in getattr(document, "tables", [])]
        pictures = getattr(document, "pictures", [])
        return ExtractionResult(
            source_id=source_id,
            source_hash=source_hash,
            extraction_id=extraction_id,
            parser=self.name,
            parser_version=version,
            markdown=markdown + ("\n" if markdown else ""),
            units=units,
            figure_count=len(pictures),
            tables=tables,
            warnings=[] if markdown else ["document_has_no_extractable_text"],
        )


def extractor_for(tool: ExtractionTool) -> Extractor:
    """Create an extractor without importing any unselected optional package."""
    implementations: dict[ExtractionTool, type[Extractor]] = {
        ExtractionTool.PYPDF: PypdfExtractor,
        ExtractionTool.PYTHON_DOCX: PythonDocxExtractor,
        ExtractionTool.MARKITDOWN: MarkItDownExtractor,
        ExtractionTool.DOCLING: DoclingExtractor,
    }
    try:
        implementation = implementations[tool]
    except KeyError as error:
        raise ValueError(f"{tool.value} is a supplementary parser, not a text extractor") from error
    return implementation()


def _package_version(distribution: str) -> str:
    try:
        return importlib.metadata.version(distribution)
    except importlib.metadata.PackageNotFoundError:
        return "unknown"
